import os
import argparse
from docx import Document
from docx.shared import Pt, Cm
import datetime

# return all lines scanned.
def process(dir, langs=[], remove_first_comment=True, remove_comment=False):
  accept_all_file = False
  if 'all' in langs:
    accept_all_file = True
  dir_len = len(dir)
  documents = [] # { title: '', lines: []}

  def parse_file(file, cstyle_comment=True):
    try:
      # is binary?
      with open(file, "r") as textlines:
        fullpath = file.path
        relapath = fullpath[dir_len:]
        lines = []
        # ingores
        copyright_removed = False
        # real_line = 0
        in_comments = False
        for line_ in textlines:
          line = line_[:-1]
          # ingore empty line
          striped_line = line.strip()
          if len(striped_line) == 0:
            continue
          # real_line += 1
          if remove_first_comment:
            if not copyright_removed:
              if cstyle_comment:
                if striped_line.startswith('//'):
                  continue
                if striped_line.startswith('/*'):
                  continue
                if striped_line.startswith('*'):
                  continue
                if '*/' in striped_line:
                  copyright_removed = True
                  continue
              else:
                if striped_line.startswith('#'):
                  continue
                if striped_line.startswith(';'):
                  continue
              copyright_removed = True
          # ingore comments
          if remove_comment:
            if cstyle_comment:
              if '*/' in striped_line:
                in_comments = False
                continue
              if in_comments:
                continue
              if '/*' in striped_line: # c-style lang
                if '*/' in striped_line:
                  # remove comment
                  comment_begin = striped_line.find('/*')
                  comment_end = striped_line.find('*/') + 2
                  if (comment_end - comment_begin <= len(striped_line)):
                    continue
                  else:
                    lines.append(line) # there exists codes except comment.
                    continue
                else:
                  in_comments = True
                continue
              if striped_line.startswith('//'): # c-style lang
                continue
            else:
              if striped_line.startswith(';'):  # python, elixir, ...
                continue
              if striped_line.startswith('#'):  # python, elixir, ...
                continue
          lines.append(line)
        print(f'  file: #lines={len(lines): >5}  {fullpath}')
        # print(f'\tline: {len(lines)}')
        documents.append((relapath, fullpath, lines))
    except UnicodeDecodeError:
      pass # Fond non-text data

  def scan_subdir(subdir):
    # print(f'\tpath: {subdir}')
    for file in os.scandir(subdir):
      if file.is_dir():
        scan_subdir(file.path)
      else:
        ext = file.name.split('.')[-1:][0]
        cstyle_comment = ext not in ['py', 'iex', 'ex']
        if accept_all_file:
          parse_file(file, cstyle_comment=cstyle_comment)
        else:
          if ext in langs:
            parse_file(file, cstyle_comment=cstyle_comment)

  scan_subdir(dir)
  return documents

def export(sourcecodes=[], filename='output.docx', params={}):
  document = Document()

  sections = document.sections
  section = sections[0]
  section.top_margin = Cm(0.8)
  section.bottom_margin = Cm(0.8)
  section.left_margin = Cm(1)
  section.right_margin = Cm(1)
  section.header.is_linked_to_previous = True
  # styles
  filenamestyle = document.styles['Header']
  font = filenamestyle.font
  font.name = 'Arial'
  font.size = Pt(12)
  font.bold = True
  filenameformat = filenamestyle.paragraph_format
  filenameformat.space_before = Pt(10)
  filenameformat.space_after = Pt(6)
  codestyle = document.styles['Normal']
  font = codestyle.font
  font.name = 'Arial'
  font.size = Pt(10)
  codeformat = codestyle.paragraph_format
  codeformat.space_before = Pt(1)
  codeformat.space_after = Pt(1)

  document.add_heading(f'{params["name"]}源代码')
  document.add_paragraph(f'Copyright © {params["company"]} {datetime.date.today().strftime("%Y")}')
  document.add_paragraph(f'Version {params["version"]}')
  total_line = 1
  for source in sourcecodes:
    paragraph = document.add_paragraph(source[0])
    paragraph.style = filenamestyle
    lines = source[2]
    for line in lines:
      code = document.add_paragraph(f'{total_line: >6}\t{line}')
      code.style = codestyle
      total_line += 1
    # paragraph.add_run()
  document.save(filename)

def main():
  parser = argparse.ArgumentParser(
    prog='china-software-copyrighter',
    description='中国特色软件著作权代码生成器',
    epilog='Copyright(youxingz) © 2023'
  )
  parser.add_argument('-d', '--dir', required=True, help='项目所在目录')
  parser.add_argument('-m', '--main', required=False, help='主函数入口所在文件')
  parser.add_argument('-o', '--out', default='output.docx', help='输出文件名称，尽量以 .docx 结尾')
  parser.add_argument('-l', '--languages', default='all', help='需要扫描的语言文件格式，用 "," 隔开')
  parser.add_argument('-n', '--name', default='', help='项目名称')
  parser.add_argument('-c', '--company', default='', help='公司/著作者名称')
  parser.add_argument('-v', '--version', default='1.0.0', help='版本号')
  parser.add_argument('-kc', '--keepcomments', action='store_true', help='保留注释')
  parser.add_argument('-kr', '--keepcopyright', action='store_true', help='保留文件顶部版权信息')
  args = parser.parse_args()
  # print(f'dir   = {args.dir}')
  # print(f'main  = {args.main}')
  # print(f'out   = {args.out}')
  # print(f'langs = {args.languages}')
  print('source code scanning...')
  docs = process(args.dir, args.languages.split(','), remove_first_comment=(not args.keepcomments), remove_comment=(not args.keepcopyright))
  # rearrange
  if len(args.main) != 0:
    main = None
    for doc in docs:
      if doc[1].endswith(args.main):
        main = doc
        docs.remove(main)
        print(f'  main file detected: {main[1]}')
        break
    docs.insert(0, main)
  print(f'{len(docs)} files selected, docx file generating...')
  # output:
  export(docs, args.out, {
    'name': args.name,
    'company': args.company,
    'version': args.version,
  })
  print(f'task done, export file: {args.out}')

if __name__ == '__main__':
  main()
